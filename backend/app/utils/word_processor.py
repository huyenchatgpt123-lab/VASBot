from docx import Document as DocxDocument
from typing import List, Dict, Any

from app.config import settings
from app.utils.chunking import chunk_text


def extract_text_from_docx(filepath: str) -> List[Dict[str, Any]]:
    """Extract text from Word document, including table rows."""
    doc = DocxDocument(filepath)
    full_text = []
    for para in doc.paragraphs:
        if para.text.strip():
            full_text.append(para.text)
    for table in doc.tables:
        for row in table.rows:
            cells = []
            for cell in row.cells:
                raw = " ".join(p.text.strip() for p in cell.paragraphs if p.text.strip())
                value = " ".join(raw.split()).strip()
                if value:
                    cells.append(value)
            if cells:
                # Keep a stable delimiter so downstream timeline parsing can map columns.
                full_text.append(" | ".join(cells))

    combined = "\n".join(full_text)
    chars_per_page = 3000
    pages = []
    for i in range(0, len(combined), chars_per_page):
        page_text = combined[i:i + chars_per_page]
        if page_text.strip():
            pages.append({"page_number": (i // chars_per_page) + 1, "text": page_text})

    return pages


def process_docx(filepath: str, document_id: int) -> tuple:
    """Extract text from docx, chunk, and return chunk metadata."""
    pages = extract_text_from_docx(filepath)
    all_chunks = []

    for page_data in pages:
        page_text = page_data["text"]
        page_number = page_data["page_number"]
        chunks = chunk_text(page_text, settings.CHUNK_SIZE, settings.CHUNK_OVERLAP)

        for chunk_content in chunks:
            all_chunks.append({
                "document_id": document_id,
                "content": chunk_content,
                "page_number": page_number,
            })

    return all_chunks, len(pages)
